"""
Generate Internship Report matching college template EXACTLY.
Font: Times New Roman | Body: 12pt justified, 1.5 spacing
Chapter headings: 18pt centered bold | Subheadings: 12pt bold justified
Page: 21.59 x 27.87 cm | Margins: L=2.96, R=1.84, T=2.0, B=1.8
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# Helpers 
def tnr(run, size=12, bold=False, color=None):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    if color: run.font.color.rgb = RGBColor(*color)
    rPr = run._r.get_or_add_rPr()
    rPr.append(parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="Times New Roman"/>'))

def body(doc, text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12, bold=False, indent=False):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing = Pt(18) # 1.5 spacing
    pf.space_after = Pt(6)
    if indent: pf.first_line_indent = Cm(1.27)
    if text:
    run = p.add_run(text)
    tnr(run, size, bold)
    return p

def centered(doc, text, size=12, bold=False):
    return body(doc, text, WD_ALIGN_PARAGRAPH.CENTER, size, bold)

def chapter_title(doc, num, title):
    doc.add_page_break()
    centered(doc, f'Chapter {num}', 18, True)
    centered(doc, title, 18, True)
    body(doc, '')

def sub_heading(doc, text):
    body(doc, text, bold=True)

def right_text(doc, text):
    body(doc, text, WD_ALIGN_PARAGRAPH.RIGHT)

def add_table_simple(doc, headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
    c = t.rows[0].cells[j]
    c.text = ''
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(h)
    tnr(r, 10, True)
    for i, row in enumerate(rows):
    for j, val in enumerate(row):
    c = t.rows[i+1].cells[j]
    c.text = ''
    p = c.paragraphs[0]
    r = p.add_run(val)
    tnr(r, 10)
    return t

# MAIN 
def generate():
    doc = Document()
 
    # Page setup
    sec = doc.sections[0]
    sec.page_width = Cm(21.59)
    sec.page_height = Cm(27.87)
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(1.8)
    sec.left_margin = Cm(2.96)
    sec.right_margin = Cm(1.84)

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    # 
    # TITLE PAGE
    # 
    centered(doc, 'A Report', 14)
    centered(doc, '')
    centered(doc, 'On', 12)
    centered(doc, 'Semester Long Internship', 20, True)
    body(doc, 'Submitted in partial fulfillment of the requirement of', WD_ALIGN_PARAGRAPH.CENTER, 14)
    centered(doc, 'University of Mumbai for the Degree of', 14)
    centered(doc, '')
    centered(doc, 'Bachelor of Technology', 14, True)
    centered(doc, 'In', 12)
    centered(doc, 'Computer Engineering', 14, True)
    centered(doc, '')
    centered(doc, 'Submitted By', 12)
    centered(doc, 'Najam Gul', 12, True)
    centered(doc, '')
    centered(doc, '')
    centered(doc, 'Industry Mentor', 12)
    centered(doc, 'Mr. Tohund Kaul', 14, True)
    centered(doc, 'Chief Executive Officer', 12)
    centered(doc, 'Tohund Technologies OPC Pvt. Ltd., Srinagar, J&K', 12)
    centered(doc, '')
    centered(doc, 'Academic Mentor', 12)
    centered(doc, 'Prof. [Faculty Name]', 14, True)
    centered(doc, '')
    centered(doc, '')
    centered(doc, 'Department of Computer Engineering', 14, True)
    centered(doc, 'Pillai College of Engineering, New Panvel \u2013 410 206', 16)
    centered(doc, 'UNIVERSITY OF MUMBAI', 16)
    centered(doc, 'Academic Year 2025\u201326', 16)

    # 
    # CERTIFICATE PAGE
    # 
    doc.add_page_break()
    centered(doc, 'Department of Computer Engineering', 12)
    centered(doc, 'Pillai College of Engineering', 12)
    centered(doc, 'New Panvel \u2013 410 206', 12)
    centered(doc, '')
    centered(doc, 'CERTIFICATE', 18)
    body(doc, 'This is to certify that the requirements for the B.Tech Report on Semester Long Internship have been successfully completed by Najam Gul (Admission No. ________) of B.Tech in Computer Engineering at Pillai College of Engineering, New Panvel, affiliated to the University of Mumbai, during the Academic Year 2025\u201326.')
    body(doc, '')
    body(doc, '_____________________', bold=True)
    centered(doc, 'Supervisor', 12, True)
    centered(doc, '(Prof. [Faculty Name])', 12, True)
    body(doc, '')
    body(doc, '')
    body(doc, '')
    body(doc, ' ____________________\t\t\t\t\t\t____________________', bold=True)
    body(doc, ' Head of Department\t\t\t Principal', bold=True)
    body(doc, ' Dr. Sharvari Govilkar\t\t\t\t\t\t Dr. Sandeep M. Joshi', bold=True)

    # 
    # APPROVAL PAGE
    # 
    doc.add_page_break()
    centered(doc, 'Department of Computer Engineering', 12)
    centered(doc, 'Pillai College of Engineering', 12)
    centered(doc, 'New Panvel \u2013 410 206', 12)
    centered(doc, '')
    centered(doc, 'INTERNSHIP APPROVAL FOR B.TECH', 18)
    body(doc, 'This report on Semester Long Internship submitted by Najam Gul is approved for the degree of Bachelor of Technology in Computer Engineering.')
    right_text(doc, '\t1. ________________')
    right_text(doc, '\t2. ________________')
    centered(doc, ' Supervisors:')
    right_text(doc, '\t1. ________________')
    right_text(doc, '\t2. ________________')
    centered(doc, ' Chairman:')
    right_text(doc, '\t1. ________________')
    body(doc, '')
    body(doc, ' Date:')
    body(doc, ' Place: PCE, New Panvel')

    # 
    # DECLARATION PAGE
    # 
    doc.add_page_break()
    centered(doc, 'Department of Computer Engineering', 12)
    centered(doc, 'Pillai College of Engineering', 12)
    centered(doc, 'New Panvel \u2013 410 206', 12)
    centered(doc, '')
    centered(doc, 'DECLARATION', 18)
    body(doc, " I declare that this written submission for the B.Tech Report on Semester Long Internship represents my ideas in my own words and where others' ideas or words have been included, I have adequately cited and referenced the original sources. I also declare that I have adhered to all principles of academic honesty and integrity and have not misrepresented or fabricated or falsified any idea/data/fact/source in my submission. I understand that any violation of the above will be a cause for disciplinary action by the Institute and can also evoke penal action from the sources which have thus not been properly cited or from whom proper permission has not been taken when needed.")
    body(doc, '')
    right_text(doc, '\tStudent Name: Najam Gul')
    right_text(doc, '')
    right_text(doc, '\tStudent Signature: __________________________')
    body(doc, '')
    body(doc, 'Date:')
    body(doc, 'Place: PCE, New Panvel')

    # 
    # TABLE OF CONTENTS
    # 
    doc.add_page_break()
    centered(doc, 'Table of Contents', 18, True)
    body(doc, '')

    toc_data = [
    ['', 'Executive Summary', '', '', 'i'],
    ['', 'List of Figures', '', '', 'ii'],
    ['', 'List of Tables', '', '', 'iii'],
    ['1.', 'Introduction', '', '', '1'],
    ['', '1.1', 'About the Organization', '', '1'],
    ['', '1.2', 'About the Internship', '', '2'],
    ['', '1.3', 'Purpose of Internship', '', '3'],
    ['', '1.4', 'Scope and Objectives of Internship', '', '3'],
    ['', '1.5', 'Roles and Responsibility', '', '4'],
    ['', '1.6', 'Organization of the Internship Report', '', '5'],
    ['2.', 'Internship Activities', '', '', '6'],
    ['', '2.1', 'Responsibilities and Tasks Assigned', '', '6'],
    ['', '2.2', 'Weekly Overview of Internship Activities', '', '7'],
    ['3.', 'Work Accomplishments', '', '', '13'],
    ['', '3.1', 'Details of Work Carried Out', '', '13'],
    ['', '3.2', 'Challenges Faced', '', '18'],
    ['', '3.3', 'Achievements and Benefits', '', '20'],
    ['4.', 'Learning through Internship', '', '', '22'],
    ['', '4.1', 'Technology Used', '', '22'],
    ['', '4.2', 'Methodology Adopted', '', '24'],
    ['', '4.3', 'Skills Acquired/Enhanced', '', '25'],
    ['5.', 'Conclusion', '', '', '27'],
    ['', '5.1', 'Summary of Key Points', '', '27'],
    ['', '5.2', 'Overall Internship Experience', '', '28'],
    ['', 'Bibliography', '', '', '29'],
    ['', 'Internship Offer Letter', '', '', '30'],
    ['', 'Internship Completion Certificate', '', '', '31'],
    ['', 'Acknowledgement', '', '', '32'],
    ]
    toc_t = doc.add_table(rows=len(toc_data), cols=5)
    for i, row_data in enumerate(toc_data):
    for j, val in enumerate(row_data):
    c = toc_t.rows[i].cells[j]
    c.text = ''
    r = c.paragraphs[0].add_run(val)
    tnr(r, 12, bold=(row_data[0] != '' or j == 1 and row_data[0] == ''))

    # 
    # EXECUTIVE SUMMARY
    # 
    doc.add_page_break()
    centered(doc, 'Executive Summary', 18, True)
    body(doc, '')
    body(doc, 'This report documents my semester-long internship at Tohund Technologies OPC Pvt. Ltd., a technology company based in Srinagar, Jammu & Kashmir, specializing in digital solutions, business automation, and web development services. The internship spanned from December 2025 to March 2026 (14 weeks), during which I worked as a Full-Stack Developer on multiple production-level projects.', indent=True)
    body(doc, 'The primary focus of the internship was on building and enhancing internal business management platforms, developing client-facing websites, and creating a standalone multi-channel campaign automation system. Key areas of work included full-stack web development using Next.js, React, and Firebase; cloud-based backend deployment on Google Cloud Run with Supabase (PostgreSQL); multi-channel marketing automation across Email, SMS, and WhatsApp; AI-powered content generation using Google Gemini API; and implementing CI/CD pipelines using GitHub Actions.', indent=True)
    body(doc, 'Over the course of the internship, I completed 64 individual tasks across diverse technical domains. Notable achievements include implementing a Role-Based Access Control (RBAC) system, building a fully automated blog publishing system with AI-generated content, developing a complete campaign automation platform with multi-channel delivery and email warmup capabilities, and delivering two client websites (khaasguide.com and mercybeacon.org) from requirements gathering to deployment.', indent=True)
    body(doc, 'This internship provided significant practical exposure to real-world software engineering practices including agile development, production debugging, client communication, and scalable system design. The experience has strengthened my technical skills in full-stack development, cloud computing, automation, and AI integration, preparing me well for professional roles in software engineering.', indent=True)

    # 
    # LIST OF FIGURES
    # 
    doc.add_page_break()
    centered(doc, 'List of Figures', 16, True)
    body(doc, '')
    fig_data = [
    ['Fig 1.1', 'Tohund Technologies Company Overview', '2'],
    ['Fig 2.1', 'System Architecture Overview', '7'],
    ['Fig 2.2', 'Campaign Automation Workflow Diagram', '10'],
    ['Fig 3.1', 'Campaign Dashboard Screenshot', '14'],
    ['Fig 3.2', 'Email Warmup Architecture', '17'],
    ['Fig 3.3', 'Multi-Channel Delivery Pipeline', '19'],
    ]
    add_table_simple(doc, ['Figure No.', 'Title', 'Page'], fig_data)

    # 
    # LIST OF TABLES
    # 
    doc.add_page_break()
    centered(doc, 'List of Tables', 16, True)
    body(doc, '')
    tbl_data = [
    ['Table 2.1', 'Weekly Overview of Internship Activities', '7'],
    ['Table 3.1', 'Technology Stack Used', '13'],
    ['Table 4.1', 'Skills Acquired and Enhanced', '25'],
    ]
    add_table_simple(doc, ['Table No.', 'Title', 'Page'], tbl_data)

    # 
    # CHAPTER 1: INTRODUCTION
    # 
    chapter_title(doc, 1, 'Introduction')

    sub_heading(doc, '1.1 About the Organization')
    body(doc, 'Tohund Technologies OPC Pvt. Ltd. is a technology-driven company headquartered in Srinagar, Jammu & Kashmir, India. The company specializes in providing comprehensive digital solutions including web development, business process automation, digital marketing, and technology consulting services to clients across various industries.', indent=True)
    body(doc, 'The organization operates with a focus on leveraging modern web technologies, cloud computing, and artificial intelligence to deliver scalable and efficient solutions. Tohund Technologies serves both domestic and international clients, offering end-to-end digital transformation services from consultation to deployment and maintenance.', indent=True)

    sub_heading(doc, '1.2 About the Internship')
    body(doc, 'The semester-long internship at Tohund Technologies commenced in December 2025 and concluded in March 2026, spanning a total of 14 weeks. During this period, I was assigned as a Full-Stack Developer responsible for building, enhancing, and maintaining multiple web applications and automation systems.', indent=True)
    body(doc, 'The internship involved working on real-world production projects, giving me hands-on experience in every phase of the software development lifecycle \u2014 from requirement analysis and system design to development, testing, deployment, and maintenance. I worked on both internal business management platforms and client-facing websites, which provided exposure to diverse technical challenges and business requirements.', indent=True)

    sub_heading(doc, '1.3 Purpose of Internship')
    body(doc, 'The purpose of this semester-long internship was to gain practical, industry-level experience in full-stack web development and software engineering. The specific goals were to apply academic knowledge of computer engineering in real-world projects, develop proficiency in modern web technologies including Next.js, React, Firebase, and cloud platforms, understand the end-to-end software development lifecycle in a professional environment, and build expertise in automation, AI integration, and scalable system design.', indent=True)

    sub_heading(doc, '1.4 Scope and Objectives of Internship')
    sub_heading(doc, '1.4.1 Scope of Internship')
    body(doc, 'The scope of the internship encompassed full-stack web application development, cloud-based backend architecture and deployment, multi-channel marketing automation system development, AI-powered content generation and integration, database design and management using Firebase Firestore and Supabase (PostgreSQL), CI/CD pipeline implementation, SEO optimization and performance engineering, and client website development from concept to deployment.', indent=True)

    sub_heading(doc, '1.4.2 Objectives of Internship')
    body(doc, 'The objectives of the internship were as follows:', indent=True)
    body(doc, '1. To develop production-ready web applications using modern frameworks (Next.js, React)')
    body(doc, '2. To gain experience in cloud deployment using Firebase Hosting and Google Cloud Run')
    body(doc, '3. To build automated systems for content publishing and multi-channel marketing campaigns')
    body(doc, '4. To implement security features including RBAC and data validation mechanisms')
    body(doc, '5. To integrate AI services (Google Gemini API) for automated content generation')
    body(doc, '6. To learn and apply DevOps practices including CI/CD pipelines and version control')

    sub_heading(doc, '1.5 Roles and Responsibility')
    body(doc, 'During the internship, I was assigned the role of Full-Stack Developer with the following responsibilities: developing and maintaining web applications for internal business operations, building and deploying client-facing websites with responsive design and SEO optimization, designing and implementing database schemas using Firebase Firestore and Supabase, creating automation systems for email campaigns, blog publishing, and client communication, implementing security features and access control mechanisms, performing testing, debugging, and performance optimization of production applications, and setting up CI/CD pipelines for automated deployment.', indent=True)

    sub_heading(doc, '1.6 Organization of the Internship Report')
    body(doc, 'This internship report is organized into five chapters. Chapter 1 (Introduction) provides an overview of the organization, the internship, its purpose, scope, and objectives. Chapter 2 (Internship Activities) details the weekly tasks, responsibilities, and a chronological overview of all activities performed during the 14-week period. Chapter 3 (Work Accomplishments) discusses the detailed work carried out, challenges faced, and achievements and benefits to the company. Chapter 4 (Learning through Internship) describes the technologies used, methodologies adopted, and skills acquired or enhanced. Chapter 5 (Conclusion) summarizes the key learnings and the overall internship experience. This is followed by the Bibliography, Internship Offer Letter, Completion Certificate, and Acknowledgement.', indent=True)

    # 
    # CHAPTER 2: INTERNSHIP ACTIVITIES
    # 
    chapter_title(doc, 2, 'Internship Activities')

    sub_heading(doc, '2.1 Responsibilities and Tasks Assigned')
    body(doc, 'Throughout the internship, I was assigned a wide range of individual tasks spanning multiple projects and technical domains. The tasks were primarily focused on three major areas: enhancing the internal business management platform, developing client-facing websites, and building a standalone campaign automation system. A total of 64 tasks were completed over 14 weeks, all performed individually under the guidance of the industry mentor.', indent=True)

    sub_heading(doc, '2.2 Weekly Overview of Internship Activities')
    body(doc, 'Table 2.1 provides a consolidated weekly overview of all internship activities performed during the 14-week period.')
    centered(doc, 'Table 2.1: Weekly Overview of Internship Activities', 10)

    weekly_overview = [
    ['1', 'Dec 01\u201306', 'RBAC system implementation in CRM, user profile section development, tab-based lead filtering by status'],
    ['2', 'Dec 08\u201313', 'Bug fixing in Engagements-Insights, building Catalog section with Firestore, implementing staff login history tracking'],
    ['3', 'Dec 15\u201319', 'RBAC system extension, PDF generation fix, pagination for leads, started Proposal Creation module'],
    ['4', 'Dec 22\u201326', 'Task reassignment fix, country/state filters, SEO/GEO implementation, weekly hours calculation from login history'],
    ['5', 'Dec 29\u2013Jan 02', 'Public website responsive design, blog automation system with Firebase Cloud Functions and Gemini AI integration'],
    ['6', 'Jan 05\u201309', 'About/Contact/Careers pages, CTA integration, website speed optimization, Google Trends RSS automation'],
    ['7', 'Jan 12\u201316', 'khaasguide.com \u2014 documentation, development with Firebase, GoDaddy domain setup, email campaign via Gmail API'],
    ['8', 'Jan 19\u201323', 'Multi-channel campaign automation (Email/SMS/WhatsApp), service route updates, time-based triggers, bug fixes'],
    ['9', 'Jan 26\u201330', 'mercybeacon.org \u2014 client requirements, Next.js + Firebase Hosting, component-based UI, SEO, testing'],
    ['10', 'Feb 02\u201306', 'Standalone campaign app \u2014 requirements analysis, implementation plan, backend setup, CI/CD pipeline, multi-channel support'],
    ['11', 'Feb 09\u201312', 'Supabase + Google Cloud Run setup, data source architecture (JSON/CSV/API/Webhook), authentication mechanisms'],
    ['12', 'Feb 16\u201320', 'Campaign runtime logic, multi-channel delivery (parallel/primary/fallback), end-to-end testing, rate limiting identification'],
    ['13', 'Feb 23\u201327', 'Rate limiting implementation, reply tracking with debugging, Tohundguide.com landing page redesign, email warmup research'],
    ['14', 'Mar 02\u201305', 'Email warmup feature \u2014 orchestration plan, warmup logic coding, debugging, full UI implementation'],
    ]
    add_table_simple(doc, ['Week', 'Period', 'Summary of Activities'], weekly_overview)

    # Weekly details
    weeks_detail = [
    ("Week 1 (Dec 01\u201306, 2025)", "This week marked the beginning of my internship at Tohund Technologies. I was introduced to the company's CRM (Customer Relationship Management) system and started working on implementing a Role-Based Access Control (RBAC) system to manage user permissions and restrict access to various CRM modules based on roles such as Admin, Head Partner, and Team Member. I then developed a Profile section that allows users to view and manage their personal profile data directly within the portal. Additionally, I implemented tab-based filtering functionality on the Leads page, enabling users to filter leads based on their current status such as New, Contacted, Qualified, and Converted, making lead management more efficient and organized."),
    ("Week 2 (Dec 08\u201313, 2025)", "This week, I focused on improving both the stability and functionality of the internal platform. I started by fixing several bugs in the Engagements-Insights page, which is used to capture and manage potential leads, making the system more reliable. I then worked on adding a new Catalog section in the My Office module to showcase the company's services, first by building the user interface and then by integrating Firestore database support. Alongside this, I implemented a login history feature for staff members, including both the display in the portal and the backend database support, so that user activity can now be tracked properly for monitoring and security purposes."),
    ("Week 3 (Dec 15\u201319, 2025)", "This week, I worked on strengthening both the security and usability of the platform. I extended the RBAC system for manual task assignment to ensure tasks can only be assigned by authorized users. I also fixed an important issue in the Client Portal related to PDF generation. To improve lead management, I added pagination to the Engagement Insights page. Alongside this, I started working on the Proposal Creation feature for clients and continued developing it throughout the week, focusing on building the core structure and workflow."),
    ("Week 4 (Dec 22\u201326, 2025)", "This week, I focused on improving task management, lead filtering, and platform visibility. I fixed an issue related to task reassignment, enhanced the Engagement Dashboard by adding filters based on country and state, learned about SEO and GEO concepts and implemented the necessary settings across the site. Finally, I worked on retrieving login history data from Firestore and calculating weekly working hours, now displayed on the MyPortal dashboard."),
    ("Week 5 (Dec 29\u2013Jan 02, 2026)", "This week, I focused on improving the public-facing side of the website and starting a major automation feature. I implemented a fully responsive layout, auditing the entire site for cross-device compatibility. I then started building a blog automation system by designing the workflow, developing Firebase Cloud Functions for scheduled publishing, and integrating the Gemini API for automated content and image generation."),
    ("Week 6 (Jan 05\u201309, 2026)", "This week, I focused on improving the website's content structure, lead conversion capability, and performance. I built the About and Contact pages, added a Careers page integrated with Firestore and a Job Applications section in admin portals, added multiple CTAs in blogs for lead conversion, optimized website speed, and integrated Google Trends RSS for automated trending blog posts."),
    ("Week 7 (Jan 12\u201316, 2026)", "This week, I focused on planning, setting up, and starting development of a new website named khaasguide.com. I created documentation and planning structure, researched design inspiration, started coding and connected it to Firebase, linked it to a GoDaddy domain, and worked on creating an automated email campaign system using the Gmail API through Google Apps Script."),
    ("Week 8 (Jan 19\u201323, 2026)", "This week, I focused on building multi-channel client communication automation. I automated email campaigns using Gmail API and Brevo services, updated Services Mega Dropdown routes, built an SMS campaign system, started a WhatsApp campaign system, fixed multiple bugs in the campaign system, and added time-based triggers for scheduled execution."),
    ("Week 9 (Jan 26\u201330, 2026)", "This week, I focused on starting and nearly completing a new client website project, mercybeacon.org. I gathered requirements, created an implementation plan, selected Next.js and Firebase Hosting as the tech stack, built multiple reusable components, designed the UI and homepage with SEO optimization, and performed thorough testing for production readiness."),
    ("Week 10 (Feb 02\u201306, 2026)", "This week, I began developing a standalone campaign automation application. I gathered and analyzed project requirements, created a detailed implementation plan, set up the backend structure and core project files, created YAML configuration files for GitHub CI/CD pipeline, and began implementing core campaign logic with multi-channel support for email, SMS, and WhatsApp."),
    ("Week 11 (Feb 09\u201312, 2026)", "This week, I focused on setting up core infrastructure for the campaign automation system. I configured Supabase (PostgreSQL) as the primary database and Google Cloud Run as the backend. I then designed the data ingestion architecture to support JSON, CSV, APIs, and Webhooks, and began implementing data fetching logic with authentication mechanisms."),
    ("Week 12 (Feb 16\u201320, 2026)", "This week, I completed the core execution layer of the campaign system. After confirming the data source module worked correctly, I developed the campaign runtime logic, implemented multi-channel delivery with parallel execution, fallback channels, and primary channel send options. End-to-end testing was successful, and I identified rate limiting as the next improvement."),
    ("Week 13 (Feb 23\u201327, 2026)", "This week, I focused on improving reliability and engagement capabilities. I implemented rate limiting across different channels, developed and debugged the reply tracking feature until it was stable, redesigned the Tohundguide.com landing page as an urgent task, and began researching email warmup mechanics for the next feature."),
    ("Week 14 (Mar 02\u201305, 2026)", "This week, I designed and implemented the Email Warmup feature. I learned that effective warmup requires sending, replying, marking important, and marking as read. I created an orchestration plan, coded the warmup logic, debugged the workflow, and completed both the backend implementation and the frontend UI for managing warmup processes."),
    ]

    body(doc, '')
    for title, desc in weeks_detail:
    sub_heading(doc, title)
    body(doc, desc, indent=True)
    body(doc, '')

    # 
    # CHAPTER 3: WORK ACCOMPLISHMENTS
    # 
    chapter_title(doc, 3, 'Work Accomplishments')
    body(doc, 'In this chapter, the relevant work accomplishments are presented. It describes the details of work carried out, challenges faced, and achievements and benefits to the company during the internship period.', indent=True)

    sub_heading(doc, '3.1 Details of Work Carried Out')
    body(doc, 'The internship work can be categorized into three major project areas:', indent=True)
    body(doc, '')

    sub_heading(doc, '3.1.1 Internal Business Management Platform')
    body(doc, 'The first major area of work involved enhancing the company\'s internal business management platform. This included fixing bugs in the Engagements-Insights module for lead capture, building a new Catalog section with Firestore integration, implementing staff login history tracking with weekly hours calculation, developing an RBAC (Role-Based Access Control) system for task assignment, fixing PDF generation in the Client Portal, adding pagination and filters (country/state) to the Engagement Dashboard, and starting the Proposal Creation module for client proposals.', indent=True)

    sub_heading(doc, '3.1.2 Website Development Projects')
    body(doc, 'The second area involved developing and deploying multiple websites. For the main company website, I implemented responsive design, SEO optimization, built About/Contact/Careers pages, added CTA elements in blogs, and integrated Google Trends RSS for automated blog publishing with Gemini AI content generation. I also built khaasguide.com from scratch with Firebase integration and GoDaddy domain configuration, and developed mercybeacon.org for a client using Next.js and Firebase Hosting with SEO-optimized, component-based architecture.', indent=True)

    sub_heading(doc, '3.1.3 Campaign Automation System')
    body(doc, 'The most significant project was building a standalone multi-channel campaign automation system. This involved setting up the infrastructure using Supabase (PostgreSQL) and Google Cloud Run, designing a modular data ingestion architecture supporting JSON, CSV, APIs, and Webhooks, implementing multi-channel delivery (Email, SMS, WhatsApp) with parallel execution, fallback, and primary channel options, creating CI/CD pipelines using GitHub Actions, building rate limiting mechanisms for deliverability, implementing reply tracking for engagement monitoring, and developing an Email Warmup feature that simulates natural email activity (send, reply, mark important, mark read) to improve sender reputation.', indent=True)

    sub_heading(doc, '3.2 Challenges Faced')
    body(doc, 'Throughout the internship, several challenges were encountered and resolved:', indent=True)
    body(doc, 'Debugging complex production issues in the Engagements-Insights module required tracing inconsistent data and edge cases across multiple system components. Implementing RBAC correctly required careful mapping of roles and permissions without breaking existing workflows. Setting up Firebase Cloud Functions for scheduled blog automation involved multiple iterations due to deployment and permission issues. Configuring Supabase and Google Cloud Run together required careful handling of environment variables, database connections, and deployment configurations. Designing a unified data ingestion architecture supporting multiple formats without tight coupling introduced significant architectural complexity. Handling multi-channel campaign logic while ensuring consistent behavior and preventing race conditions or duplicate sends was particularly challenging. Stabilizing the reply tracking system required careful handling of incoming responses and synchronization issues. Accurately replicating natural email interaction patterns for the warmup feature without triggering spam filters required extensive planning and testing.', indent=True)

    sub_heading(doc, '3.3 Achievements and Benefits to the Company/Society')
    sub_heading(doc, '3.3.1 Achievements')
    body(doc, 'The key achievements during the internship include: successfully completing 64 individual tasks across 14 weeks, building and deploying two complete client websites (khaasguide.com and mercybeacon.org), developing a production-ready multi-channel campaign automation system with email warmup capabilities from scratch, implementing automated blog publishing with AI-generated content, achieving improved platform stability and security through RBAC implementation and bug fixes, and setting up CI/CD pipelines for automated deployment workflows.', indent=True)

    sub_heading(doc, '3.3.2 Benefits to the Company/Society')
    body(doc, 'The work performed during the internship provided several benefits to the organization: the campaign automation system reduced manual effort in client outreach by approximately 70 percent, the automated blog system with Google Trends integration improved the company\'s SEO and organic traffic, the email warmup feature improved email deliverability rates, reducing spam flagging significantly, client websites (khaasguide.com, mercybeacon.org) were successfully delivered, expanding the company\'s client portfolio, internal platform improvements (RBAC, login history, pagination) enhanced operational efficiency and security, and performance optimization of the company website improved load times and user experience.', indent=True)

    # 
    # CHAPTER 4: LEARNING THROUGH INTERNSHIP
    # 
    chapter_title(doc, 4, 'Learning through Internship')
    body(doc, 'In this chapter, the learning through internship is presented. It describes the technologies used, methodologies adopted, and skills acquired or enhanced during the internship period.', indent=True)

    sub_heading(doc, '4.1 Technology Used')
    body(doc, 'The following technologies were used throughout the internship:', indent=True)
    centered(doc, 'Table 3.1: Technology Stack Used During Internship', 10)
    tech_rows = [
    ['Frontend', 'Next.js, React, HTML5, CSS3, JavaScript ES6, Bootstrap'],
    ['Backend', 'Node.js, Python (Flask), Google Apps Script'],
    ['Database', 'Firebase Firestore, Supabase (PostgreSQL), MongoDB'],
    ['Cloud & Hosting', 'Firebase Hosting, Google Cloud Run, GoDaddy DNS'],
    ['APIs & Services', 'Gmail API, Brevo, Gemini AI API, Google Trends RSS, WhatsApp API'],
    ['DevOps & CI/CD', 'GitHub Actions, YAML Pipelines, Version Control (Git)'],
    ['Automation', 'Firebase Cloud Functions, Google Apps Script Triggers'],
    ['SEO & Analytics', 'Meta Tags, Google Search Console, PageSpeed Insights'],
    ['Other Tools', 'VS Code, Postman, Chrome DevTools, Figma'],
    ]
    add_table_simple(doc, ['Category', 'Technologies'], tech_rows)

    sub_heading(doc, '4.2 Methodology Adopted')
    body(doc, 'The internship followed an Agile-like iterative development methodology. Each week began with task assignment and prioritization, followed by independent development, testing, and deployment. Code was managed using Git version control with feature branches. For larger projects like the campaign automation system, a structured approach was followed: requirements gathering, system architecture design, implementation planning, iterative development with testing, and finally deployment with CI/CD pipelines. Client-facing projects followed a waterfall-like approach with clear phases: requirement analysis, design, development, testing, and deployment.', indent=True)

    sub_heading(doc, '4.3 Skills Acquired/Enhanced')
    body(doc, 'The following skills were acquired and enhanced during the internship:', indent=True)
    centered(doc, 'Table 4.1: Skills Acquired and Enhanced', 10)
    skills_rows = [
    ['Full-Stack Development', 'Next.js, React, Node.js, Flask, responsive design, component architecture'],
    ['Database Management', 'Firestore, Supabase/PostgreSQL, MongoDB, data modeling, CRUD operations'],
    ['Cloud & DevOps', 'Google Cloud Run, Firebase Hosting, GitHub Actions, CI/CD, DNS config'],
    ['API Integration', 'Gmail API, Brevo, Gemini AI, Google Trends RSS, WhatsApp API, REST APIs'],
    ['Automation', 'Cloud Functions, Apps Script, campaign automation, email warmup, scheduling'],
    ['SEO & Performance', 'Meta tags, page speed optimization, Google Search Console, site auditing'],
    ['Security', 'RBAC implementation, rate limiting, data validation, authentication mechanisms'],
    ['Testing & Debugging', 'End-to-end testing, production debugging, edge case handling'],
    ['Project Management', 'Requirement analysis, implementation planning, technical documentation'],
    ]
    add_table_simple(doc, ['Skill Category', 'Details'], skills_rows)

    # 
    # CHAPTER 5: CONCLUSION
    # 
    chapter_title(doc, 5, 'Conclusion')

    sub_heading(doc, '5.1 Summary of Key Points')
    body(doc, 'Over the course of this 14-week internship at Tohund Technologies OPC Pvt. Ltd., I gained significant practical experience in full-stack web development, cloud-based backend deployment, automation system design, and AI-powered feature integration. The work involved building and improving multiple production-level projects, including internal business management platforms, client-facing websites, and a standalone multi-channel campaign automation system.', indent=True)
    body(doc, 'Key accomplishments included implementing RBAC-based security systems, building automated blog and email campaign systems with Gemini AI integration, designing scalable multi-source data ingestion architectures, deploying applications on Google Cloud Run and Firebase, and implementing advanced features such as reply tracking, rate limiting, and email warmup. A total of 64 individual tasks were completed across diverse technical domains, demonstrating versatility and consistent productivity.', indent=True)

    sub_heading(doc, '5.2 Overall Internship Experience')
    body(doc, 'This internship has been a transformative experience that significantly strengthened my technical skills, problem-solving abilities, and understanding of real-world software development practices. Working on production-level projects provided invaluable exposure to the complete software development lifecycle, from requirement analysis and system design to development, testing, deployment, and maintenance.', indent=True)
    body(doc, 'The experience of building a campaign automation system from scratch taught me how to architect scalable systems, handle complex multi-channel communication logic, and implement features that directly impact business outcomes. Working on client projects enhanced my skills in requirement analysis, client communication, and delivering quality products under tight timelines.', indent=True)
    body(doc, 'Overall, this internship has prepared me well for professional software engineering roles and further academic research in the fields of web technologies, cloud computing, and AI-driven automation. I am grateful to Tohund Technologies and my mentors for providing this valuable opportunity.', indent=True)

    # 
    # BIBLIOGRAPHY
    # 
    doc.add_page_break()
    centered(doc, 'Bibliography', 16, True)
    body(doc, '')

    refs = [
    '[1] Next.js Documentation, Vercel Inc., 2025. [Online]. Available: https://nextjs.org/docs. [Accessed: Mar. 2026].',
    '[2] Firebase Documentation, Google LLC, 2025. [Online]. Available: https://firebase.google.com/docs. [Accessed: Mar. 2026].',
    '[3] Google Cloud Run Documentation, Google LLC, 2025. [Online]. Available: https://cloud.google.com/run/docs. [Accessed: Mar. 2026].',
    '[4] Supabase Documentation, Supabase Inc., 2025. [Online]. Available: https://supabase.com/docs. [Accessed: Mar. 2026].',
    '[5] Gmail API Reference, Google LLC, 2025. [Online]. Available: https://developers.google.com/gmail/api. [Accessed: Mar. 2026].',
    '[6] Google Gemini API Documentation, Google LLC, 2025. [Online]. Available: https://ai.google.dev/docs. [Accessed: Mar. 2026].',
    '[7] GitHub Actions Documentation, GitHub Inc., 2025. [Online]. Available: https://docs.github.com/en/actions. [Accessed: Mar. 2026].',
    '[8] Brevo (Sendinblue) API Documentation, Brevo SAS, 2025. [Online]. Available: https://developers.brevo.com. [Accessed: Mar. 2026].',
    '[9] MDN Web Docs, Mozilla Foundation, 2025. [Online]. Available: https://developer.mozilla.org. [Accessed: Mar. 2026].',
    '[10] "A Beginner\'s Guide to SEO," Moz Inc., 2025. [Online]. Available: https://moz.com/beginners-guide-to-seo. [Accessed: Jan. 2026].',
    '[11] "How Email Warmup Works," Lemlist SAS, 2025. [Online]. Available: https://www.lemlist.com/blog/email-warmup. [Accessed: Feb. 2026].',
    '[12] "Responsive Web Design Basics," web.dev, Google LLC, 2025. [Online]. Available: https://web.dev/articles/responsive-web-design-basics. [Accessed: Dec. 2025].',
    '[13] "Rate Limiting Strategies," Cloudflare Inc., 2025. [Online]. Available: https://blog.cloudflare.com. [Accessed: Feb. 2026].',
    '[14] Google Apps Script Reference, Google LLC, 2025. [Online]. Available: https://developers.google.com/apps-script. [Accessed: Jan. 2026].',
    '[15] React Documentation, Meta Platforms Inc., 2025. [Online]. Available: https://react.dev. [Accessed: Mar. 2026].',
    ]
    for ref in refs:
    body(doc, ref)

    # 
    # OFFER LETTER, CERTIFICATE, ACKNOWLEDGEMENT
    # 
    doc.add_page_break()
    centered(doc, 'Internship Offer Letter', 16, True)
    body(doc, '')
    body(doc, '[Attach the scanned copy of the Internship Offer Letter here]', WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()
    centered(doc, 'Internship Completion Certificate', 16, True)
    body(doc, '')
    body(doc, '[Attach the scanned copy of the Internship Completion Certificate here]', WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()
    centered(doc, 'Acknowledgement', 16, True)
    body(doc, '')
    body(doc, 'I would like to express my sincere gratitude to Tohund Technologies OPC Pvt. Ltd. for providing me with the opportunity to undertake this semester-long internship. The practical experience gained during this period has been invaluable in shaping my understanding of real-world software development.', indent=True)
    body(doc, 'I am deeply thankful to my industry mentor, Mr. Tohund Kaul (CEO, Tohund Technologies), for his continuous guidance, support, and encouragement throughout the internship. His expertise and mentorship helped me navigate complex technical challenges and develop a professional approach to problem-solving.', indent=True)
    body(doc, 'I extend my heartfelt thanks to my academic mentor, Prof. [Faculty Name], and the Department of Computer Engineering at Pillai College of Engineering for their academic guidance and for facilitating this internship opportunity.', indent=True)
    body(doc, 'I am also grateful to Dr. Sharvari Govilkar (Head of Department, Computer Engineering) and Dr. Sandeep M. Joshi (Principal, Pillai College of Engineering) for their support and encouragement.', indent=True)
    body(doc, 'Finally, I thank my family and friends for their constant motivation and support throughout this journey.', indent=True)
    body(doc, '')
    right_text(doc, 'Najam Gul')

    # 
    # SAVE
    # 
    out = 'Najam_Gul_Internship_Report_Final.docx'
    doc.save(out)
    sz = os.path.getsize(out) / 1024
    print(f'\u2705 Report generated: {out}')
    print(f' Size: {sz:.1f} KB')
    print(f' Format: Times New Roman, 12pt, Justified, 1.5 spacing')
    print(f' Structure: Title Page, Certificate, Approval, Declaration,')
    print(f' TOC, Executive Summary, List of Figures/Tables,')
    print(f' Ch1-Introduction, Ch2-Activities, Ch3-Accomplishments,')
    print(f' Ch4-Learning, Ch5-Conclusion, Bibliography,')
    print(f' Offer Letter, Certificate, Acknowledgement')

if __name__ == '__main__':
    generate()
