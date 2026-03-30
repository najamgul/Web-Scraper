"""
Generate UPDATED IPR Documents in DOCX format with:
- All authors: Najam Gul, Kalyani, Megha, Dnyanesh
- Mentor: Ms. Neha Ashok
- Architecture diagrams and screenshots from report/PPT
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), 'IPR')
IMG_DIR = os.path.join(OUTPUT_DIR, 'images')

# ============================================================
# Helper functions
# ============================================================
def add_formatted_paragraph(doc, text, bold=False, italic=False, size=12, alignment=None, space_after=6, space_before=0, font_name='Times New Roman'):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = font_name
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    return p

def add_blank_line(doc, count=1):
    for _ in range(count):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)

def add_field_line(doc, label, value="", dotted=True):
    p = doc.add_paragraph()
    run = p.add_run(label)
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    if value:
        run2 = p.add_run(f" {value}")
        run2.font.size = Pt(12)
        run2.font.name = 'Times New Roman'
    elif dotted:
        run2 = p.add_run(" " + "." * 60)
        run2.font.size = Pt(12)
        run2.font.name = 'Times New Roman'
    p.paragraph_format.space_after = Pt(4)
    return p

def add_image_with_caption(doc, image_path, caption, width=Inches(5.5)):
    """Add an image with a centered caption below it."""
    if os.path.exists(image_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(image_path, width=width)
        p.paragraph_format.space_after = Pt(4)
        
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(caption)
        run.bold = True
        run.italic = True
        run.font.size = Pt(10)
        run.font.name = 'Times New Roman'
        cap.paragraph_format.space_after = Pt(12)
    else:
        add_formatted_paragraph(doc, f'[Image not found: {os.path.basename(image_path)}]', italic=True, size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)


# ============================================================
# DOCUMENT 1: Authorisation / No Objection Certificate (UPDATED)
# ============================================================
def generate_authorisation_letter():
    doc = Document()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
    
    # Date
    add_formatted_paragraph(doc, '27th March 2026', size=12, alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_after=18)
    
    # College address
    add_formatted_paragraph(doc, 'Pillai College of Engineering', bold=True, size=13, space_after=2)
    add_formatted_paragraph(doc, 'Dr. K. M. Vasudevan Pillai Campus, Plot No. 10, Sector 16,', size=12, space_after=2)
    add_formatted_paragraph(doc, 'New Panvel East, Navi Mumbai, Maharashtra', size=12, space_after=2)
    add_formatted_paragraph(doc, '410206', size=12, space_after=18)
    
    # Subject
    p = doc.add_paragraph()
    run = p.add_run('Subject: ')
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run2 = p.add_run('No Objection Certificate for Copyright Application')
    run2.bold = True
    run2.underline = True
    run2.font.size = Pt(12)
    run2.font.name = 'Times New Roman'
    p.paragraph_format.space_after = Pt(12)
    
    add_formatted_paragraph(doc, 'Sir/Madam,', size=12, space_after=12)
    
    # Body paragraphs - with ALL AUTHORS
    add_formatted_paragraph(doc, 
        'This is to inform you that Mr. Najam Gul, Ms. Kalyani, Ms. Megha, and Mr. Dnyanesh '
        'are students at Pillai College of Engineering. '
        'Mr. Najam Gul (applicant and author), Ms. Kalyani (author), Ms. Megha (author), '
        'and Mr. Dnyanesh (author), along with Ms. Neha Ashok (mentor and guide), are applying for copyright '
        'of a software work titled "Web Scraper for Data Extraction and Threat Intelligence: '
        'An AI-Powered Cybersecurity Analysis Tool".',
        size=12, space_after=12)
    
    add_formatted_paragraph(doc,
        'On behalf of the Principal, Pillai College of Engineering, we authorize Mr. Najam Gul '
        '(applicant and author), Ms. Kalyani (author), Ms. Megha (author), Mr. Dnyanesh (author), '
        'and Ms. Neha Ashok (mentor and guide) to apply for copyright for the above '
        'said software work and I have no objection to the aforesaid work being '
        'registered in the name of Mr. Najam Gul (applicant and author), Ms. Kalyani (author), '
        'Ms. Megha (author), Mr. Dnyanesh (author), and Ms. Neha Ashok '
        '(mentor and guide) under the Copyright Act.',
        size=12, space_after=12)
    
    add_formatted_paragraph(doc,
        'If you have any queries regarding this matter, please feel free to contact us via email '
        'principalpce@mes.ac.in or through phone.',
        size=12, space_after=18)
    
    add_formatted_paragraph(doc, 'Sincerely,', size=12, space_after=24)
    add_blank_line(doc, 2)
    
    add_formatted_paragraph(doc, 'Dr. Sandeep Joshi', bold=True, size=12, space_after=2)
    add_formatted_paragraph(doc, 'Principal', size=12, space_after=2)
    add_formatted_paragraph(doc, 'Pillai College of Engineering', size=12, space_after=12)
    
    filepath = os.path.join(OUTPUT_DIR, 'Authorisation_Letter_NOC.docx')
    doc.save(filepath)
    print(f"✅ Created: {filepath}")


# ============================================================
# DOCUMENT 2: Copyright Work Description Document (UPDATED with diagrams)
# ============================================================
def generate_work_description():
    doc = Document()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
    
    # Title page / Header
    add_formatted_paragraph(doc, 'Diary Number: _______________/2026-CO/SW', bold=True, size=12, space_after=18)
    
    add_formatted_paragraph(doc, 
        'Web Scraper for Data Extraction and Threat Intelligence:\nAn AI-Powered Cybersecurity Analysis Tool',
        bold=True, size=16, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
    
    # ALL AUTHORS
    add_formatted_paragraph(doc, 'Authors:', bold=True, size=13, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    add_formatted_paragraph(doc, 'Najam Gul (Applicant & Author)', size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_formatted_paragraph(doc, 'Kalyani (Author)', size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_formatted_paragraph(doc, 'Megha (Author)', size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_formatted_paragraph(doc, 'Dnyanesh (Author)', size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    add_formatted_paragraph(doc, 'Mentor & Guide: Ms. Neha Ashok', bold=True, size=13, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    add_formatted_paragraph(doc, 'Pillai College of Engineering', italic=True, size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
    
    # ---- Section 1: Introduction ----
    add_formatted_paragraph(doc, '1. Introduction', bold=True, size=14, space_after=6, space_before=12)
    
    add_formatted_paragraph(doc, 
        'Web scraping and threat intelligence are critical components of modern cybersecurity. '
        'This work presents a comprehensive web-based platform that combines automated data extraction '
        'with AI-powered threat intelligence analysis. The tool enables security professionals to '
        'perform deep analysis of Indicators of Compromise (IOCs), conduct bulk scanning operations, '
        'and generate detailed threat intelligence reports.',
        size=12, space_after=8)
    
    add_formatted_paragraph(doc,
        'The platform integrates multiple threat intelligence APIs and data sources, including '
        'VirusTotal, AbuseIPDB, Shodan, AlienVault OTX, and others, to provide a unified view '
        'of cyber threats. It employs machine learning models for automated threat classification '
        'and risk scoring, making it a valuable tool for Security Operations Centers (SOC) and '
        'cybersecurity researchers.',
        size=12, space_after=12)
    
    # Landing Page Screenshot
    add_image_with_caption(doc, 
        os.path.join(IMG_DIR, 'docx_b_image_007.png'),
        'Fig 1: Landing Page - Threat Intelligent Web Scraper',
        width=Inches(5.0))
    
    # ---- Section 2: System Architecture ----
    add_formatted_paragraph(doc, '2. System Architecture', bold=True, size=14, space_after=6, space_before=12)
    
    add_formatted_paragraph(doc,
        'The platform follows a layered architecture pattern with clear separation of concerns. '
        'The system consists of a Flask-based backend, a responsive web frontend, multiple API '
        'integration modules, a machine learning classification engine, and an AI-powered analysis layer.',
        size=12, space_after=8)
    
    # Architecture Flowchart
    add_image_with_caption(doc,
        os.path.join(IMG_DIR, 'docx_b_image_002.png'),
        'Fig 2: System Architecture Flowchart',
        width=Inches(4.5))
    
    add_formatted_paragraph(doc,
        'The architecture encompasses the following key components:', bold=True, size=12, space_after=6)
    
    arch_components = [
        ('Frontend Layer', 'HTML5/CSS3/JavaScript with Bootstrap for responsive design, Chart.js for data visualization, and dynamic AJAX-based interactions.'),
        ('Application Layer', 'Flask web framework with SQLAlchemy ORM, Flask-Login for authentication, and Flask-Migrate for database migrations.'),
        ('API Integration Layer', 'Parallel API calls to VirusTotal, Shodan, AbuseIPDB, AlienVault OTX, and Google Custom Search Engine for comprehensive threat data collection.'),
        ('ML Classification Engine', 'Random Forest classifier with TF-IDF vectorization for automated threat categorization into malware, phishing, botnet, ransomware, and benign classes.'),
        ('AI Analysis Layer', 'Google Gemini integration for advanced natural language threat analysis and report generation.'),
        ('Data Storage Layer', 'SQLite/PostgreSQL database for scan history, user management, and investigation notebooks.')
    ]
    
    for title, desc in arch_components:
        p = doc.add_paragraph()
        run = p.add_run(f'• {title}: ')
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        run2 = p.add_run(desc)
        run2.font.size = Pt(12)
        run2.font.name = 'Times New Roman'
        p.paragraph_format.space_after = Pt(6)
    
    add_blank_line(doc)
    
    # Use Case Diagram
    add_image_with_caption(doc,
        os.path.join(IMG_DIR, 'docx_b_image_014.png'),
        'Fig 3: Use Case Diagram - Web Scraper for Data Extraction and Threat Intelligence',
        width=Inches(5.5))
    
    # ---- Section 3: Threat Intelligence ----
    doc.add_page_break()
    add_formatted_paragraph(doc, '3. Threat Intelligence', bold=True, size=14, space_after=6, space_before=12)
    
    add_formatted_paragraph(doc,
        'Threat Intelligence refers to the collection, analysis, and dissemination of information '
        'about potential or current cyber threats targeting an organization. It helps security teams '
        'make informed decisions about defending against cyber attacks by providing context about '
        'threat actors, their tactics, techniques, and procedures (TTPs), and indicators of compromise (IOCs).',
        size=12, space_after=8)
    
    add_formatted_paragraph(doc, 'Types of Threat Intelligence:', bold=True, size=12, space_after=6)
    
    ti_types = [
        ('Strategic Threat Intelligence', 'High-level information about the threat landscape, trends, and risks that helps executives and decision-makers understand the broader cybersecurity context and allocate resources effectively.'),
        ('Tactical Threat Intelligence', 'Information about the TTPs used by threat actors, helping security teams understand how attacks are carried out and develop appropriate defenses and detection mechanisms.'),
        ('Operational Threat Intelligence', 'Details about specific attacks or campaigns, including the who, what, when, and how of cyber threats, enabling incident response teams to take immediate action.'),
        ('Technical Threat Intelligence', 'Specific technical indicators such as IP addresses, domain names, file hashes, and URLs associated with malicious activity, used for automated detection and blocking.')
    ]
    
    for title, desc in ti_types:
        p = doc.add_paragraph()
        run = p.add_run(f'{title}: ')
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        run2 = p.add_run(desc)
        run2.font.size = Pt(12)
        run2.font.name = 'Times New Roman'
        p.paragraph_format.space_after = Pt(6)
    
    # ---- Section 4: Indicators of Compromise ----
    add_formatted_paragraph(doc, '4. Indicators of Compromise (IOCs)', bold=True, size=14, space_after=6, space_before=12)
    
    add_formatted_paragraph(doc,
        'Indicators of Compromise are artifacts observed on a network or operating system that indicate '
        'a potential intrusion or malicious activity. The platform supports analysis of the following IOC types:',
        size=12, space_after=8)
    
    ioc_types = [
        'IP Addresses — IPv4 and IPv6 addresses associated with malicious activity',
        'Domain Names — Domains linked to phishing, malware distribution, or command & control servers',
        'URLs — Specific web addresses known to host malware or phishing pages',
        'File Hashes — MD5, SHA-1, and SHA-256 hashes of known malicious files',
        'Email Addresses — Email accounts used in phishing or spam campaigns',
        'CVE Identifiers — Common Vulnerabilities and Exposures identifiers for known security flaws'
    ]
    
    for ioc in ioc_types:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(ioc)
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
    
    # IOC Scanner Screenshot
    add_blank_line(doc)
    add_image_with_caption(doc,
        os.path.join(IMG_DIR, 'docx_b_image_004.png'),
        'Fig 4: IOC Scanner Interface - Input Data Page',
        width=Inches(5.0))
    
    # ---- Section 5: Web Scraping ----
    add_formatted_paragraph(doc, '5. Web Scraping for Threat Data', bold=True, size=14, space_after=6, space_before=12)
    
    add_formatted_paragraph(doc,
        'Web scraping is an automated technique used to extract data from websites. In the context of '
        'cybersecurity, web scraping enables the collection of threat data from multiple online sources, '
        'including threat intelligence feeds, security blogs, vulnerability databases, and dark web forums.',
        size=12, space_after=8)
    
    add_formatted_paragraph(doc,
        'The platform employs sophisticated scraping techniques to gather threat data from diverse sources, '
        'normalize it into a consistent format, and correlate findings across multiple data points to provide '
        'comprehensive threat assessments.',
        size=12, space_after=12)
    
    # ---- Section 6: Machine Learning ----
    doc.add_page_break()
    add_formatted_paragraph(doc, '6. Machine Learning for Threat Classification', bold=True, size=14, space_after=6, space_before=12)
    
    add_formatted_paragraph(doc,
        'The platform integrates machine learning models for automated threat classification. '
        'A Random Forest classifier is trained on labeled threat data to categorize IOCs into threat '
        'categories such as malware, phishing, botnet, ransomware, and benign. The model achieves '
        'high accuracy through feature engineering that incorporates both technical indicators and '
        'contextual information from multiple threat intelligence sources.',
        size=12, space_after=8)
    
    add_formatted_paragraph(doc, 'Key ML features include:', bold=True, size=12, space_after=6)
    
    ml_features = [
        'Multi-model ensemble approach (Random Forest, Gradient Boosting, SVM, Neural Networks)',
        'TF-IDF vectorization for text-based threat indicators',
        'Feature extraction from API responses across multiple threat intelligence platforms',
        'Real-time classification with confidence scoring',
        'Continuous model improvement through feedback loops'
    ]
    
    for feat in ml_features:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(feat)
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
    
    add_blank_line(doc)
    
    # Confusion Matrix
    add_image_with_caption(doc,
        os.path.join(IMG_DIR, 'all_models_confusion_matrix.png'),
        'Fig 5: ML Model Confusion Matrix - Multi-Model Classification Results',
        width=Inches(4.5))
    
    # Scan Results - AI Analysis
    add_image_with_caption(doc,
        os.path.join(IMG_DIR, 'docx_b_image_012.png'),
        'Fig 6: Scan Results Page - AI Threat Analysis and Classification',
        width=Inches(5.0))
    
    # ---- Section 7: Methodology ----
    doc.add_page_break()
    add_formatted_paragraph(doc, '7. Methodology', bold=True, size=14, space_after=6, space_before=12)
    
    add_formatted_paragraph(doc, 'The development of this platform followed a systematic methodology:', size=12, space_after=8)
    
    methodology_steps = [
        ('Data Collection & API Integration', 
         'Integration with multiple threat intelligence APIs (VirusTotal, AbuseIPDB, Shodan, AlienVault OTX, URLScan.io, GreyNoise, etc.) to collect comprehensive threat data for each IOC submitted for analysis.'),
        ('Data Processing & Normalization',
         'Raw data from various APIs is processed, cleaned, and normalized into a consistent format. This includes extracting relevant features, handling missing data, and creating unified threat profiles.'),
        ('Machine Learning Model Training',
         'Training of classification models using labeled threat data. The models are evaluated using cross-validation, confusion matrices, and standard metrics (accuracy, precision, recall, F1-score).'),
        ('Web Application Development',
         'Development of a Flask-based web application with a professional SOC-style interface. The application provides individual IOC analysis, bulk scanning, investigation notebooks, and report generation.'),
        ('Threat Scoring & Classification',
         'Implementation of a multi-dimensional threat scoring system that combines API-based risk indicators with ML-based classifications to provide comprehensive threat assessments.'),
        ('Report Generation',
         'Automated generation of detailed threat intelligence reports in multiple formats, including on-screen reports and downloadable summaries (CSV, JSON, Excel, PDF) for offline analysis and sharing.')
    ]
    
    for i, (title, desc) in enumerate(methodology_steps, 1):
        p = doc.add_paragraph()
        run = p.add_run(f'{i}. {title}: ')
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        run2 = p.add_run(desc)
        run2.font.size = Pt(12)
        run2.font.name = 'Times New Roman'
        p.paragraph_format.space_after = Pt(8)
    
    # Pipeline Execution Time
    add_blank_line(doc)
    add_image_with_caption(doc,
        os.path.join(IMG_DIR, 'fig_pipeline_execution.png'),
        'Fig 7: Pipeline Component Execution Time Analysis (n=50 queries)',
        width=Inches(5.0))
    
    # ---- Section 8: Platform Features & Screenshots ----
    doc.add_page_break()
    add_formatted_paragraph(doc, '8. Platform Features', bold=True, size=14, space_after=6, space_before=12)
    
    features = [
        ('Single IOC Analysis', 'Deep analysis of individual indicators with data from multiple threat intelligence sources, ML-based classification, and comprehensive threat reports.'),
        ('Bulk IOC Scanner', 'Upload and analyze multiple IOCs simultaneously with parallel processing, progress tracking, and consolidated results.'),
        ('Investigation Notebooks', 'Create and manage investigation cases, associate multiple IOCs and findings, add notes and timelines, and generate case reports.'),
        ('Dashboard & Analytics', 'Real-time dashboard showing scan statistics, threat distribution, recent activity, and trending threats.'),
        ('API Integrations', 'Seamless integration with VirusTotal, AbuseIPDB, Shodan, AlienVault OTX, URLScan.io, GreyNoise, IPQualityScore, and more.'),
        ('Report Generation', 'Automated generation of detailed threat intelligence reports with visualizations, risk scores, and actionable recommendations.')
    ]
    
    for title, desc in features:
        p = doc.add_paragraph()
        run = p.add_run(f'• {title}: ')
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        run2 = p.add_run(desc)
        run2.font.size = Pt(12)
        run2.font.name = 'Times New Roman'
        p.paragraph_format.space_after = Pt(6)
    
    add_blank_line(doc)
    
    # Dashboard Screenshot
    add_image_with_caption(doc,
        os.path.join(IMG_DIR, 'docx_b_image_001.png'),
        'Fig 8: Threat Intelligence Dashboard - Analytics Overview',
        width=Inches(5.5))
    
    # Dashboard Analytics (more details)
    add_image_with_caption(doc,
        os.path.join(IMG_DIR, 'docx_b_image_011.png'),
        'Fig 9: Dashboard Analytics - Top Malicious IPs, Geographic Distribution & Live Threat Feed',
        width=Inches(5.5))
    
    # Threat Analysis Results
    doc.add_page_break()
    add_image_with_caption(doc,
        os.path.join(IMG_DIR, 'docx_b_image_008.png'),
        'Fig 10: Threat Analysis - Key Threat Indicators and Recommended Action',
        width=Inches(5.0))
    
    # AlienVault OTX Results
    add_image_with_caption(doc,
        os.path.join(IMG_DIR, 'docx_b_image_009.png'),
        'Fig 11: AlienVault OTX Threat Intelligence - Threat Score and Recent Threat Pulses',
        width=Inches(5.0))
    
    # Classification & Export
    add_image_with_caption(doc,
        os.path.join(IMG_DIR, 'docx_b_image_015.png'),
        'Fig 12: Classification Overview, VirusTotal Detections, and Export Results',
        width=Inches(5.5))
    
    # ---- Section 9: Technology Stack ----
    doc.add_page_break()
    add_formatted_paragraph(doc, '9. Technology Stack', bold=True, size=14, space_after=6, space_before=12)
    
    tech_stack = [
        ('Backend', 'Python, Flask, SQLAlchemy, Flask-Migrate, Flask-Login'),
        ('Frontend', 'HTML5, CSS3, JavaScript, Bootstrap, Chart.js'),
        ('Machine Learning', 'scikit-learn, Random Forest, Gradient Boosting, TF-IDF Vectorizer'),
        ('AI/LLM', 'Google Gemini API for advanced threat analysis'),
        ('Database', 'SQLite (development), PostgreSQL (production)'),
        ('APIs', 'VirusTotal, AbuseIPDB, Shodan, AlienVault OTX, URLScan.io, GreyNoise, IPQualityScore, Google CSE'),
        ('Deployment', 'Vercel (frontend), Gunicorn (WSGI server)')
    ]
    
    table = doc.add_table(rows=len(tech_stack)+1, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    hdr[0].text = 'Component'
    hdr[1].text = 'Technologies'
    for p in hdr[0].paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(12)
    for p in hdr[1].paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(12)
    for i, (comp, tech) in enumerate(tech_stack, 1):
        table.rows[i].cells[0].text = comp
        table.rows[i].cells[1].text = tech
        for cell in table.rows[i].cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(11)
    
    add_blank_line(doc)
    
    # Latency Comparison
    add_image_with_caption(doc,
        os.path.join(IMG_DIR, 'fig_latency_comparison.png'),
        'Fig 13: API Latency Comparison across Threat Intelligence Sources',
        width=Inches(4.5))
    
    # ---- Section 10: Applications ----
    add_formatted_paragraph(doc, '10. Applications', bold=True, size=14, space_after=6, space_before=12)
    
    add_formatted_paragraph(doc,
        'The platform has various impactful applications across cybersecurity domains:',
        size=12, space_after=8)
    
    applications = [
        ('Security Operations Centers (SOC)', 'Analysts can use the platform for rapid IOC triage, threat scoring, and investigation management, reducing response times and improving threat detection accuracy.'),
        ('Incident Response', 'During security incidents, the platform enables rapid analysis of suspicious indicators, helping responders quickly determine the nature and scope of threats.'),
        ('Threat Hunting', 'Proactive threat hunters can leverage the bulk scanning feature to analyze large sets of potentially malicious indicators and identify emerging threats.'),
        ('Vulnerability Management', 'CVE lookup and analysis features help organizations assess the impact of known vulnerabilities on their infrastructure.'),
        ('Compliance & Reporting', 'Automated report generation provides documentation suitable for compliance requirements and stakeholder communication.'),
        ('Academic Research', 'The platform serves as a practical tool for cybersecurity education and research, demonstrating real-world applications of ML in security.')
    ]
    
    for title, desc in applications:
        p = doc.add_paragraph()
        run = p.add_run(f'• {title}: ')
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        run2 = p.add_run(desc)
        run2.font.size = Pt(12)
        run2.font.name = 'Times New Roman'
        p.paragraph_format.space_after = Pt(8)
    
    # ---- Section 11: Conclusion ----
    doc.add_page_break()
    add_formatted_paragraph(doc, '11. Conclusion', bold=True, size=14, space_after=6, space_before=12)
    
    add_formatted_paragraph(doc,
        'This work presents a comprehensive web-based threat intelligence platform that combines '
        'web scraping, API integrations, and machine learning to provide automated cybersecurity analysis. '
        'The platform offers a professional SOC-style interface for analyzing indicators of compromise, '
        'conducting bulk scans, managing investigations, and generating detailed threat intelligence reports.',
        size=12, space_after=8)
    
    add_formatted_paragraph(doc,
        'The machine learning models integrated into the platform demonstrate the effectiveness of '
        'AI-powered threat classification, achieving high accuracy in categorizing threats across multiple '
        'categories. The multi-source API integration provides comprehensive threat coverage, while the '
        'investigation notebook feature enables structured case management.',
        size=12, space_after=8)
    
    add_formatted_paragraph(doc,
        'The platform serves as a valuable resource for cybersecurity professionals, SOC analysts, '
        'and researchers, providing practical tools for threat intelligence analysis and demonstrating '
        'the application of modern web technologies and machine learning in cybersecurity.',
        size=12, space_after=12)
    
    # AI Analysis Wireframe
    add_image_with_caption(doc,
        os.path.join(IMG_DIR, 'pptx_image_056.png'),
        'Fig 14: Platform UI Wireframe - Scan Results, AI Analysis, and Confidence Metrics',
        width=Inches(4.0))
    
    # ---- Section 12: Author Information ----
    add_formatted_paragraph(doc, '12. Author Information', bold=True, size=14, space_after=6, space_before=12)
    
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    author_data = [
        ('Applicant & Author', 'Najam Gul'),
        ('Author', 'Kalyani'),
        ('Author', 'Megha'),
        ('Author', 'Dnyanesh'),
        ('Mentor & Guide', 'Ms. Neha Ashok'),
        ('Institution', 'Pillai College of Engineering'),
    ]
    
    for i, (role, name) in enumerate(author_data):
        table.rows[i].cells[0].text = role
        table.rows[i].cells[1].text = name
        for cell in table.rows[i].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(12)
                    run.font.name = 'Times New Roman'
        # Bold the role column
        for paragraph in table.rows[i].cells[0].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    add_blank_line(doc)
    
    # Software availability
    add_formatted_paragraph(doc, 'Software Availability:', bold=True, size=12, space_after=6)
    add_formatted_paragraph(doc, 'The source code and documentation are available at the project repository on GitHub.', size=12, space_after=12)
    
    filepath = os.path.join(OUTPUT_DIR, 'Copyright_Work_Description.docx')
    doc.save(filepath)
    print(f"✅ Created: {filepath}")


# ============================================================
# Main
# ============================================================
if __name__ == '__main__':
    print("Generating UPDATED IPR Documents...")
    print(f"Output directory: {OUTPUT_DIR}")
    print()
    generate_authorisation_letter()
    generate_work_description()
    print()
    print("✅ All updated IPR documents generated successfully!")
    print("📁 Files are in:", OUTPUT_DIR)
