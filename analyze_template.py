from docx import Document

doc = Document(r'Najam Gul - Internship Report Doc.docx')

with open('template_analysis.txt', 'w', encoding='utf-8') as f:
    section = doc.sections[0]
    f.write('=== PAGE SETUP ===\n')
    f.write(f'Page width: {section.page_width.cm:.2f} cm\n')
    f.write(f'Page height: {section.page_height.cm:.2f} cm\n')
    f.write(f'Top margin: {section.top_margin.cm:.2f} cm\n')
    f.write(f'Bottom margin: {section.bottom_margin.cm:.2f} cm\n')
    f.write(f'Left margin: {section.left_margin.cm:.2f} cm\n')
    f.write(f'Right margin: {section.right_margin.cm:.2f} cm\n\n')

    f.write('=== ALL PARAGRAPHS ===\n')
    for i, p in enumerate(doc.paragraphs):
    style = p.style.name if p.style else 'None'
    alignment = str(p.alignment) if p.alignment else 'None'
    text = p.text[:150] if p.text else '[empty]'
 
    font_info = ''
    for run in p.runs[:2]:
    fn = run.font
    sz = str(fn.size.pt) if fn.size else 'inh'
    nm = fn.name or 'inh'
    cl = str(fn.color.rgb) if fn.color and fn.color.rgb else 'inh'
    bd = run.bold
    font_info += f' [bold={bd}, size={sz}, font={nm}, color={cl}]'
 
    f.write(f'P{i:3d} | {style:25s} | {alignment:35s} | {text}\n')
    if font_info:
    f.write(f' {font_info}\n')

    f.write('\n=== TABLES ===\n')
    for i, table in enumerate(doc.tables):
    f.write(f'\nTable {i}: {len(table.rows)} rows x {len(table.columns)} cols\n')
    for j, row in enumerate(table.rows):
    cells = []
    for cell in row.cells:
    ct = cell.text[:50].replace('\n', ' ')
    cells.append(ct)
    f.write(f' Row {j}: {cells}\n')

print('Saved to template_analysis.txt')
