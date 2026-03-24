"""
Convert PRESENTATION_PREPARATION_GUIDE.md to a professionally formatted .docx file.
"""

import re
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

def set_cell_shading(cell, color_hex):
    """Set background color of a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_styled_paragraph(doc, text, style='Normal', bold=False, italic=False, 
    font_size=None, color=None, alignment=None, space_after=None, space_before=None):
    """Add a paragraph with custom styling."""
    p = doc.add_paragraph()
    if style and style != 'Normal':
    try:
    p.style = style
    except:
    pass
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if font_size:
    run.font.size = Pt(font_size)
    if color:
    run.font.color.rgb = RGBColor(*color)
    if alignment is not None:
    p.alignment = alignment
    if space_after is not None:
    p.paragraph_format.space_after = Pt(space_after)
    if space_before is not None:
    p.paragraph_format.space_before = Pt(space_before)
    return p

def add_table_from_rows(doc, headers, rows, header_color="1F4E79", stripe=True):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(header.strip())
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(255, 255, 255)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_shading(cell, header_color)

    # Data rows
    for r_idx, row in enumerate(rows):
    for c_idx, cell_text in enumerate(row):
    cell = table.rows[r_idx + 1].cells[c_idx]
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(cell_text.strip())
    run.font.size = Pt(9)
    if stripe and r_idx % 2 == 1:
    set_cell_shading(cell, "E8F0FE")

    return table

def parse_md_table(lines):
    """Parse markdown table lines into headers and rows."""
    headers = []
    rows = []
    for i, line in enumerate(lines):
    line = line.strip()
    if not line.startswith('|'):
    continue
    cells = [c.strip() for c in line.split('|')[1:-1]]
    if i == 0 or not headers:
    if not all(set(c) <= set('- :') for c in cells):
    headers = cells
    elif all(set(c) <= set('- :') for c in cells):
    continue # separator line
    else:
    rows.append(cells)
    return headers, rows

def clean_emoji(text):
    """Remove emoji characters but keep the text."""
    # Keep common emoji-like markers but clean up for docx
    return text

def process_inline_formatting(paragraph, text):
    """Process bold, italic, and code inline formatting."""
    # Pattern to match **bold**, *italic*, `code`, and plain text
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`[^`]+`)', text)
    for part in parts:
    if not part:
    continue
    if part.startswith('**') and part.endswith('**'):
    run = paragraph.add_run(part[2:-2])
    run.bold = True
    run.font.size = Pt(10)
    elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
    run = paragraph.add_run(part[1:-1])
    run.italic = True
    run.font.size = Pt(10)
    elif part.startswith('`') and part.endswith('`'):
    run = paragraph.add_run(part[1:-1])
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)
    else:
    run = paragraph.add_run(part)
    run.font.size = Pt(10)

def create_docx():
    md_path = r"PRESENTATION_PREPARATION_GUIDE.md"
    docx_path = r"PRESENTATION_PREPARATION_GUIDE.docx"
 
    with open(md_path, 'r', encoding='utf-8') as f:
    content = f.read()
 
    lines = content.split('\n')
 
    doc = Document()
 
    # ----- Page Setup -----
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
 
    # ----- Default Font -----
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)
 
    # ----- Heading Styles -----
    for level in range(1, 5):
    heading_style = doc.styles[f'Heading {level}']
    heading_style.font.name = 'Calibri'
    heading_style.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    if level == 1:
    heading_style.font.size = Pt(22)
    heading_style.font.bold = True
    elif level == 2:
    heading_style.font.size = Pt(16)
    heading_style.font.bold = True
    elif level == 3:
    heading_style.font.size = Pt(13)
    heading_style.font.bold = True
    elif level == 4:
    heading_style.font.size = Pt(11)
    heading_style.font.bold = True

    # TITLE PAGE
    for _ in range(4):
    doc.add_paragraph()
 
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(" Presentation & Report\nPreparation Guide")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
 
    doc.add_paragraph()
 
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle_p.add_run("Web Scraper for Data Extraction\nand Threat Intelligence")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
 
    doc.add_paragraph()
 
    details = [
    "B.Tech Project-C | Academic Year 2025–26",
    "Pillai College of Engineering (Autonomous), New Panvel",
    "University of Mumbai · Department of Computer Engineering",
    ]
    for detail in details:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(detail)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x59, 0x56, 0x59)
 
    doc.add_paragraph()
    doc.add_paragraph()
 
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run("March 2026")
    run.font.size = Pt(12)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
 
    doc.add_page_break()

    # TABLE OF CONTENTS placeholder
    doc.add_heading("Table of Contents", level=1)
    toc_items = [
    "1. Purpose of This Document",
    "2. Team Members & Speaking Distribution",
    "3. Project Overview",
    "4. Slide-by-Slide Explanation",
    " 4.1 Slide 1–3: Title, Team & Roadmap",
    " 4.2 Slide 4: Abstract",
    " 4.3 Slide 5: Introduction & Fundamentals",
    " 4.4 Slide 6: Project Objectives",
    " 4.5 Slide 7: Literature Survey",
    " 4.6 Slide 8: Research Gaps",
    " 4.7 Slide 9: Existing vs. Proposed System",
    " 4.8 Slide 10: System Architecture",
    " 4.9 Slide 11: IOC Detection & API Routing",
    " 4.10 Slide 12: Parallel Orchestration Engine",
    " 4.11 Slide 13: ML Classification Pipeline",
    " 4.12 Slide 14: LLM-Powered Threat Enrichment",
    " 4.13 Slide 15: Technology Stack",
    " 4.14 Slide 16: GUI & Interface Walkthrough",
    " 4.15 Slide 17: Results & AI-Generated Analysis",
    " 4.16 Slide 18: Performance Evaluation",
    " 4.17 Slide 19: Conclusion & Future Scope",
    " 4.18 Slide 20: Acknowledgements & References",
    "5. Expected Viva Questions & Answers (20 Questions)",
    "6. Key Terms Glossary",
    "7. Presentation Tips",
    ]
    for item in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(item)
    run.font.size = Pt(10)
    if not item.startswith(" "):
    run.bold = True
 
    doc.add_page_break()

    # Now parse the actual markdown content
    i = 0
    in_code_block = False
    code_lines = []
    in_table = False
    table_lines = []
    in_blockquote = False
 
    while i < len(lines):
    line = lines[i]
    stripped = line.strip()
 
    # --- Code blocks ---
    if stripped.startswith('```'):
    if in_code_block:
    # End code block
    code_text = '\n'.join(code_lines)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    # Add light gray background shading
    pPr = p._p.get_or_add_pPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0F0F0" w:val="clear"/>')
    pPr.append(shading)
    code_lines = []
    in_code_block = False
    else:
    # Start code block - flush any pending table
    if in_table and table_lines:
    headers, rows = parse_md_table(table_lines)
    if headers and rows:
    add_table_from_rows(doc, headers, rows)
    doc.add_paragraph()
    table_lines = []
    in_table = False
    in_code_block = True
    i += 1
    continue
 
    if in_code_block:
    code_lines.append(line.rstrip())
    i += 1
    continue
 
    # --- Tables ---
    if stripped.startswith('|') and '|' in stripped[1:]:
    if not in_table:
    in_table = True
    table_lines = []
    table_lines.append(stripped)
    i += 1
    continue
    else:
    if in_table and table_lines:
    headers, rows = parse_md_table(table_lines)
    if headers and rows:
    add_table_from_rows(doc, headers, rows)
    doc.add_paragraph()
    elif headers:
    add_table_from_rows(doc, headers, [])
    doc.add_paragraph()
    table_lines = []
    in_table = False
 
    # --- Empty lines ---
    if not stripped:
    i += 1
    continue
 
    # --- Horizontal rules ---
    if stripped in ['---', '***', '___']:
    # Add a thin horizontal line
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
    f'<w:pBdr {nsdecls("w")}>'
    f'<w:bottom w:val="single" w:sz="4" w:space="1" w:color="CCCCCC"/>'
    f'</w:pBdr>'
    )
    pPr.append(pBdr)
    i += 1
    continue
 
    # --- Headings ---
    heading_match = re.match(r'^(#{1,4})\s+(.+)$', stripped)
    if heading_match:
    level = len(heading_match.group(1))
    heading_text = heading_match.group(2).strip()
    # Clean emoji from headings for cleaner docx
    h = doc.add_heading(heading_text, level=level)
    i += 1
    continue
 
    # --- Blockquotes ---
    if stripped.startswith('>'):
    quote_text = stripped[1:].strip()
    # Remove any leading >
    quote_text = re.sub(r'^>\s*', '', quote_text)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    # Add left border
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
    f'<w:pBdr {nsdecls("w")}>'
    f'<w:left w:val="single" w:sz="12" w:space="4" w:color="2E75B6"/>'
    f'</w:pBdr>'
    )
    pPr.append(pBdr)
    process_inline_formatting(p, quote_text)
    # Color the text
    for run in p.runs:
    run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    run.italic = True
    i += 1
    continue
 
    # --- Unordered list items ---
    list_match = re.match(r'^(\s*)[-*]\s+(.+)$', stripped)
    if list_match:
    indent_level = len(line) - len(line.lstrip())
    list_text = list_match.group(2)
    p = doc.add_paragraph()
    p.style = 'List Bullet'
    if indent_level > 0:
    p.paragraph_format.left_indent = Cm(1.5 + (indent_level * 0.3))
    process_inline_formatting(p, list_text)
    i += 1
    continue
 
    # --- Ordered list items ---
    olist_match = re.match(r'^(\s*)\d+[\.\)]\s+(.+)$', stripped)
    if olist_match:
    list_text = olist_match.group(2)
    p = doc.add_paragraph()
    p.style = 'List Number'
    process_inline_formatting(p, list_text)
    i += 1
    continue
 
    # --- Normal paragraph ---
    p = doc.add_paragraph()
    process_inline_formatting(p, stripped)
    i += 1
 
    # Flush any remaining table
    if in_table and table_lines:
    headers, rows = parse_md_table(table_lines)
    if headers and rows:
    add_table_from_rows(doc, headers, rows)
 
    # Save
    doc.save(docx_path)
    print(f"Successfully created: {docx_path}")
    print(f"File size: {os.path.getsize(docx_path) / 1024:.1f} KB")

if __name__ == "__main__":
    create_docx()
